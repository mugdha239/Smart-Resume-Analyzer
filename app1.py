from flask import Flask, request, redirect, url_for, flash, render_template, send_file
import os
from werkzeug.utils import secure_filename
import pdfplumber
import docx
import re
from nltk.tokenize import word_tokenize

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['EXTRACTED_FILES_FOLDER'] = 'extracted_files/'
app.config['EXTRACTED_SKILLS_FOLDER'] = 'extracted_skills/'  # New folder for storing skills
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'doc', 'docx'}
app.secret_key = 'your_secret_key'

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# def extract_text_from_pdf(filepath):
#     text = ''
#     try:
#         with pdfplumber.open(filepath) as pdf:
#             for page in pdf.pages:
#                 text += page.extract_text() + '\n'
#     except Exception as e:
#         print(f"Error processing PDF file: {e}")
#     return text

# def extract_text_from_docx(filepath):
#     text = ''
#     try:
#         doc = docx.Document(filepath)
#         for para in doc.paragraphs:
#             text += para.text + '\n'
#     except Exception as e:
#         print(f"Error processing DOCX file: {e}")
#     return text

# def extract_skills(text):
#     pattern = re.compile(r"(skills|technical skills|key skills|competencies|expertise)([\s\S]*?)(experience|education|projects|certifications|$)", re.IGNORECASE)
#     match = pattern.search(text)
    
#     if match:
#         skills_section = match.group(2).strip()  
       
#         skills_list = [skill.strip() for skill in re.split(r'[,\n]+', skills_section) if skill.strip()]
#         return ', '.join(skills_list)
#     return None


def extract_text_from_pdf(filepath):
    text = ''
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                # Extract text and identify columns if possible
                text += page.extract_text() + '\n'
    except Exception as e:
        print(f"Error processing PDF file: {e}")
    return text

def extract_text_from_docx(filepath):
    text = ''
    try:
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            text += para.text + '\n'
    except Exception as e:
        print(f"Error processing DOCX file: {e}")
    return text

def extract_skills(text):
    # Refined pattern to capture various skills sections
    pattern = re.compile(r"(skills|technical skills|key skills|competencies|expertise)([\s\S]*?)(experience|education|projects|certifications|$)", re.IGNORECASE)
    match = pattern.search(text)
    
    if match:
        skills_section = match.group(2).strip()  # Extract the matched skills section

        # Further processing to handle column-based formats
        # Normalize spaces and line breaks
        skills_section = re.sub(r'\s+', ' ', skills_section)

        # Split skills by common delimiters and join with commas
        skills_list = [skill.strip() for skill in re.split(r'[,\n]+', skills_section) if skill.strip()]
        
        # Join skills with commas
        return ', '.join(skills_list)
    
    return None




@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Process the file based on its type
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(filepath)
            elif filename.lower().endswith('.docx'):
                text = extract_text_from_docx(filepath)

            # Save extracted text to a .txt file
            text_filename = filename.rsplit('.', 1)[0] + '.txt'
            text_filepath = os.path.join(app.config['EXTRACTED_FILES_FOLDER'], text_filename)
            with open(text_filepath, 'w') as text_file:
                text_file.write(text)

            # Extract skills and save separately
            skills = extract_skills(text)
            if skills:
                skills_filename = filename.rsplit('.', 1)[0] + '_skills.txt'
                skills_filepath = os.path.join(app.config['EXTRACTED_SKILLS_FOLDER'], skills_filename)
                with open(skills_filepath, 'w') as skills_file:
                    skills_file.write(skills)

            # Provide a success message
            flash('File uploaded and processed successfully!')
            return render_template('upload.html', filename=filename, text_filename=text_filename, skills_filename=skills_filename, message='File uploaded successfully!')

        else:
            flash('This format is not allowed. Only DOC and PDF are allowed.')
            return redirect(request.url)

    return render_template('upload.html')

@app.route('/download/<filename>')
def download_file(filename):
    return send_file(os.path.join(app.config['EXTRACTED_FILES_FOLDER'], filename), as_attachment=True)

if __name__ == "__main__":
    # Ensure the upload, extracted files, and extracted skills directories exist
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    if not os.path.exists(app.config['EXTRACTED_FILES_FOLDER']):
        os.makedirs(app.config['EXTRACTED_FILES_FOLDER'])
    if not os.path.exists(app.config['EXTRACTED_SKILLS_FOLDER']):
        os.makedirs(app.config['EXTRACTED_SKILLS_FOLDER'])
    
    app.run(debug=True)
